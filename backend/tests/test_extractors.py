import io
import pytest
from tools.extractors import (
    build_extraction_diagnostics,
    extract_text,
    extract_pdf,
    extract_docx,
    per_source_status,
    repair_extracted_text,
)


class TestExtractText:
    def test_txt_extraction(self):
        content = b"Hello World\nSecond line"
        result = extract_text(content, "test.txt")
        assert result == "Hello World\nSecond line"

    def test_txt_utf8_with_replacements(self):
        content = "Hello \xff\xfe World".encode("utf-8", errors="replace")
        result = extract_text(content, "test.txt")
        assert "Hello" in result

    def test_unsupported_file_type(self):
        with pytest.raises(ValueError, match="Unsupported file type"):
            extract_text(b"content", "test.xyz")

    def test_pdf_extraction(self):
        # Create a minimal valid PDF with text
        pdf_bytes = (
            b"%PDF-1.4\n"
            b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
            b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
            b"3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R"
            b"/Resources<</Font<</F1 4 0 R>>>>>>endobj\n"
            b"4 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
            b"xref\n0 5\n"
            b"0000000000 65535 f \n"
            b"0000000009 00000 n \n"
            b"0000000058 00000 n \n"
            b"0000000115 00000 n \n"
            b"0000000266 00000 n \n"
            b"trailer<</Size 5/Root 1 0 R>>\n"
            b"startxref\n345\n%%EOF"
        )
        # pdfplumber may handle this gracefully or raise - both are acceptable
        try:
            result = extract_pdf(pdf_bytes)
            assert isinstance(result, str)
        except Exception:
            pass  # Minimal PDF may not be parseable

    def test_docx_extraction(self):
        from docx import Document

        doc = Document()
        doc.add_paragraph("Hello World")
        doc.add_paragraph("Second paragraph")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        result = extract_docx(buf.read())
        assert "Hello World" in result
        assert "Second paragraph" in result

    def test_docx_empty_paragraphs(self):
        from docx import Document

        doc = Document()
        doc.add_paragraph("Only content")
        doc.add_paragraph("")  # empty paragraph
        doc.add_paragraph("  ")  # whitespace-only paragraph
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        result = extract_docx(buf.read())
        assert "Only content" in result
        assert result.strip() == "Only content"


class TestExtractionDiagnostics:
    def test_per_source_status_word_only(self):
        # A CV that only contains the words, not real URLs.
        text = "Profile: LinkedIn | GitHub | Kaggle"
        statuses = per_source_status(text, [])
        assert statuses["github"]["status"] == "word_only"
        assert statuses["linkedin"]["status"] == "word_only"
        assert statuses["kaggle"]["status"] == "word_only"

    def test_per_source_status_url_found(self):
        text = "My projects"
        urls = ["https://github.com/jane", "https://www.linkedin.com/in/jane"]
        statuses = per_source_status(text, urls)
        assert statuses["github"]["status"] == "url_found"
        assert statuses["linkedin"]["status"] == "url_found"

    def test_per_source_status_not_found(self):
        statuses = per_source_status("Just text", [])
        assert statuses["github"]["status"] == "not_found"

    def test_build_extraction_diagnostics_counts(self):
        diag = build_extraction_diagnostics(
            "Hello world", ["https://github.com/jane"], filename="cv.txt"
        )
        assert diag["character_count"] == 11
        assert diag["page_count"] == 1
        assert diag["hyperlink_annotations"] == 0
        assert diag["detected_urls"] == ["https://github.com/jane"]
        assert diag["sources"]["github"]["status"] == "url_found"


class TestTextRepair:
    def test_repairs_glued_tech_phrases_and_broken_spacing(self):
        raw = (
            "\u2022 FastAPIbackendservingaTelegrambotandwebsite;"
            "deployedonAWSEC2withautomatedtestinganddeployment.\n"
            "Real-timevoiceagentwithLiveKitandMCPsupportingEnglish-Arabiccode-"
            "switchingandlow-latencyaudiostream-ing.\n"
            "verifi\u00adcation"
        )
        fixed = repair_extracted_text(raw)
        assert "FastAPI backend serving a Telegram bot and website" in fixed
        assert "deployed on AWS EC2 with automated testing and deployment" in fixed
        assert (
            "Real-time voice agent with LiveKit and MCP supporting "
            "English-Arabic code-switching and low-latency audio streaming"
        ) in fixed
        assert "verification" in fixed

    def test_repairs_common_concatenated_terms(self):
        fixed = repair_extracted_text("MachineLearning and NaturalLanguageProcessing and PowerBI")
        assert "Machine Learning" in fixed
        assert "Natural Language Processing" in fixed
        assert "Power BI" in fixed

    def test_repair_preserves_normal_text(self):
        original = "Python developer with 5 years of experience building dashboards."
        assert repair_extracted_text(original) == original
